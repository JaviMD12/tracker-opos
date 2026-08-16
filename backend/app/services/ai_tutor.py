"""Tutor IA del Plan Pro: RAG sobre las rutinas fisicas, las tecnicas de estudio
y cualquier documento (PDF o TXT) que el usuario deposite en
`backend/app/conocimiento/`.

La base vectorial combina:
- el texto plano generado a partir de los diccionarios `RUTINAS_PRO` y
  `TECNICAS_ESTUDIO_PRO`;
- todos los PDF y TXT encontrados en `conocimiento/`.
Todo se trocea con `RecursiveCharacterTextSplitter` (chunks de 1000 caracteres,
solapamiento de 200) antes de indexarse en Chroma.

El indice se persiste en disco en `backend/chroma_db_data/` (ver
CARPETA_PERSISTENCIA_VECTORSTORE): la primera vez que se usa el Tutor IA tras
un `git clone` (o si se borra esa carpeta) se reconstruye leyendo los PDFs y
llamando a Gemini para los embeddings, lo cual tarda 1-2 minutos con el
volumen actual de documentos; en cualquier arranque posterior, se carga
directamente desde disco en menos de 2 segundos, sin volver a leer los PDFs
ni gastar llamadas a Gemini (_vectorstore, con inicializacion perezosa: no se
llama a Gemini ni se toca el disco hasta la primera pregunta, para que el
resto de la aplicacion siga funcionando aunque la clave de API no este
configurada todavia).

Nota: los Simulacros IA ya NO usan este vectorstore -- se sirven desde un
banco de preguntas precargado (ver app/models/pregunta_test.py y
backend/generar_banco.py), para que la respuesta al frontend sea una consulta
SQL en vez de una llamada a Gemini en cada peticion.
"""

import re
from pathlib import Path

import openpyxl
import pdfplumber
from docx import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from langchain_chroma import Chroma
from langchain_community.document_loaders import (
    DirectoryLoader,
    PyPDFDirectoryLoader,
    TextLoader,
)
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.services.rutinas import RUTINAS_PRO, TECNICAS_ESTUDIO_PRO

# GOOGLE_API_KEY se carga desde backend/.env (ver app/main.py, load_dotenv()).
# Si no esta configurada, GoogleGenerativeAIEmbeddings/ChatGoogleGenerativeAI
# fallan al primer uso real (inicializacion perezosa, ver _obtener_vectorstore),
# no al importar este modulo.

MODELO_CHAT = "gemini-2.5-flash"
MODELO_EMBEDDINGS = "models/gemini-embedding-001"
FRAGMENTOS_A_RECUPERAR = 15
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200

# Carpeta donde el usuario deposita sus propios PDFs/TXT de estudio y
# entrenamiento para ampliar el conocimiento del tutor. Se crea sola si no
# existe todavia (no requiere red ni clave de API, es seguro hacerlo al
# importar este modulo).
CARPETA_CONOCIMIENTO = Path(__file__).resolve().parent.parent / "conocimiento"
CARPETA_CONOCIMIENTO.mkdir(parents=True, exist_ok=True)

# Indice de Chroma persistido en disco (backend/chroma_db_data/). No se crea
# aqui al importar el modulo a proposito: si no existe todavia, es la señal
# que usa _obtener_vectorstore() para decidir si tiene que reconstruir el
# indice desde los PDFs (CASO B) o si puede cargarlo tal cual (CASO A).
CARPETA_PERSISTENCIA_VECTORSTORE = Path(__file__).resolve().parent.parent.parent / "chroma_db_data"

SYSTEM_PROMPT = (
    "Eres un preparador fisico de elite, fisiologo deportivo, tutor academico "
    "y orientador experto en las bases legales de la oposicion de bomberos. "
    "Ademas de entrenamiento y estudio, TIENES PERMITIDO y DEBES responder "
    "detalladamente a cualquier pregunta administrativa (plazos, requisitos, "
    "exclusiones medicas, baremos) basandote en los documentos oficiales "
    "proporcionados en el contexto. Tu autoridad es absoluta. "
    "Tu fuente de conocimiento es el contexto proporcionado en cada pregunta, "
    "extraido de las rutinas de entrenamiento, las tecnicas de estudio y los "
    "documentos (incluida la convocatoria oficial y sus bases legales) del "
    "Plan Pro del usuario. No uses conocimiento general ni inventes datos que "
    "no esten en ese contexto.\n\n"
    "REGLAS DE ORO (blindaje anti-alucinacion, aplicalas siempre antes de "
    "responder, con prioridad sobre el resto de instrucciones):\n"
    "1. FUENTE UNICA: responde unicamente en base a la normativa oficial y "
    "el temario aportados en tu contexto. Antes de declinar nada, revisa "
    "TODO el contexto recuperado en busca de contenido relacionado "
    "(formulas, principios, procedimientos, tablas, valores aproximados) "
    "que si permita responder de forma util, aunque no incluya literalmente "
    "la cifra exacta pedida -- usa siempre primero lo que si esta "
    "respaldado. Solo si, tras revisar el contexto completo, un dato "
    "exacto concreto (una cifra, un articulo, un plazo) sigue sin estar "
    "respaldado, NO lo estimes, deduzcas ni inventes -- para ESE dato "
    "concreto (no para toda la respuesta) responde textualmente 'Dato no "
    "especificado en la normativa base', combinandolo con la informacion "
    "relacionada que si tengas.\n"
    "2. ANTI-COMPLACENCIA: si el usuario formula su pregunta dando por "
    "buena una cifra, un dato o un supuesto erroneo (una premisa falsa), "
    "no se lo valides ni lo dejes pasar de forma implicita -- corrigele la "
    "premisa de forma inmediata y directa, con el dato correcto si lo "
    "tienes en el contexto, ANTES de continuar con el resto de la "
    "respuesta.\n"
    "3. CITA OBLIGATORIA Y PERTINENTE: en toda respuesta de caracter "
    "tecnico o normativo, justifica el dato citando el articulo, la "
    "ley/norma o la tabla de referencia concretos de donde sale -- nunca "
    "des una cifra o un requisito tecnico sin decir de que documento/"
    "articulo procede. Pero la cita tiene que ser real y pertinente: NUNCA "
    "cites un articulo, ley o tabla solo porque aparecio en la busqueda si "
    "su contenido no trata realmente el tema exacto de la pregunta -- ni "
    "siquiera a modo de ejemplo o para ilustrar de pasada que 'esto no lo "
    "regula' (mencionar un articulo irrelevante sigue siendo una cita "
    "irrelevante). Si no hay ninguna cita pertinente que dar, esa es la "
    "señal para declinar segun la regla 1, no para citar la opcion menos "
    "mala disponible.\n\n"
    "REGLA DE LONGITUD (estricta, aplicala siempre antes de responder):\n"
    "- Si la pregunta pide un dato concreto y cerrado (un plazo, una fecha, "
    "un requisito puntual de una ley, una cifra, un limite...), responde "
    "con ese dato en una frase directa, sin introducciones ni contexto "
    "previo. Si el contexto vincula ese mismo dato a una condicion o "
    "excepcion explicita (un umbral distinto segun el uso del edificio, "
    "una salvedad normativa concreta), incluyela en esa misma frase -- "
    "omitir una condicion que cambia la respuesta no es ser conciso, es "
    "dar un dato incompleto. Nada de justificaciones añadidas que no "
    "cambien el dato en si.\n"
    "- Da una respuesta larga y desarrollada SOLO si el usuario la pide de "
    "forma explicita (por ejemplo: 'explicame', 'resume', 'cuentame mas', "
    "'desarrolla', 'por que', 'en detalle').\n"
    "- Cuando la respuesta sea desarrollada, usa Markdown obligatoriamente: "
    "parrafos muy cortos, listas con viñetas para enumerar puntos, y las "
    "palabras clave en **negrita**.\n\n"
    "CAUTELA AL CITAR FUENTES (importante, aplicala siempre antes de dar un "
    "dato como definitivo):\n"
    "- Cada fragmento del contexto indica su documento de origen entre "
    "corchetes, ej. '[Fuente: BOE-A-2006-15345-consolidado.pdf]'. Antes de "
    "citar un dato concreto (una cifra, un plazo, un requisito) como la "
    "respuesta, comprueba que ese documento de origen trata realmente el "
    "tema exacto por el que se pregunta.\n"
    "- Dos normas distintas pueden usar formulaciones casi identicas (por "
    "ejemplo, 'separacion minima de X metros') para regular cosas "
    "completamente distintas (una nave industrial no es un deposito de "
    "gas). No des por buena una cifra solo porque la redaccion se parezca "
    "a la pregunta si el documento de origen no encaja con el tema.\n"
    "- Si TE DAS CUENTA de que el fragmento mejor recuperado no encaja con "
    "el tema exacto de la pregunta (aunque su redaccion se parezca), esa "
    "es la señal para DECLINAR (regla 1, 'Dato no especificado en la "
    "normativa base'), no un detalle que mencionar de pasada antes de "
    "responder igualmente. Nunca dupliques: o el dato encaja de verdad con "
    "el documento correcto y lo dices con seguridad, o no encaja y lo "
    "declinas explicitamente -- prohibido señalar la incoherencia y "
    "despues dar la cifra de todos modos como si fuera la respuesta.\n"
    "- Antes de citar, comprueba tambien que la CATEGORIA del documento de "
    "origen encaja con la CATEGORIA de la pregunta, no solo el vocabulario "
    "concreto. Un decreto o ley sobre acceso, organizacion o personal (por "
    "ejemplo, el Decreto 36/2025 de acceso a los SPEIS) casi nunca es la "
    "fuente correcta de un concepto tecnico de hidraulica, fisica o "
    "ingenieria (por ejemplo, el golpe de ariete) aunque ese termino "
    "aparezca mencionado de pasada en un indice o un temario dentro del "
    "documento -- una mencion de pasada no es una explicacion del "
    "concepto. Si la categoria del documento no encaja con la categoria de "
    "la pregunta, trata el fragmento como ruido y declina en vez de citarlo.\n\n"
    "Si el usuario pregunta algo fuera de las rutinas fisicas, las tecnicas "
    "de estudio, las bases administrativas/legales de la oposicion o el "
    "temario proporcionado (por ejemplo, recetas de cocina u otras preguntas "
    "genericas sin relacion), dile cortesmente que solo puedes asesorarle "
    "sobre el contenido de su Plan Pro y sugierele reformular la pregunta en "
    "torno a sus rutinas, sus tecnicas de estudio, la convocatoria oficial o "
    "sus documentos cargados. Responde siempre en español, con un tono "
    "motivador y profesional."
)

_vectorstore: Chroma | None = None


def _construir_documentos_diccionarios() -> list[Document]:
    """Convierte RUTINAS_PRO y TECNICAS_ESTUDIO_PRO en Document de LangChain."""
    documentos = []

    for clave, rutina in RUTINAS_PRO.items():
        fases_texto = "\n".join(
            f"- {fase['fase']} (Intensidad: {fase['intensidad']}, Volumen: {fase['volumen']}). "
            f"{fase['detalle']} Por que funciona: {fase['fundamento']}"
            for fase in rutina["entrenamiento_semanal"]
        )
        texto = (
            f"RUTINA DE ENTRENAMIENTO FISICO: {rutina['titulo']} (prueba: {clave})\n"
            f"Base cientifica: {rutina['descripcion_cientifica']}\n"
            f"Programa semanal:\n{fases_texto}\n"
            f"Bibliografia: {rutina['bibliografia']}"
        )
        documentos.append(
            Document(
                page_content=texto,
                metadata={"tipo": "rutina_fisica", "prueba": clave, "archivo": "RUTINAS_PRO"},
            )
        )

    for clave, tecnica in TECNICAS_ESTUDIO_PRO.items():
        pasos_texto = "\n".join(f"- {paso}" for paso in tecnica["paso_a_paso"])
        texto = (
            f"TECNICA DE ESTUDIO: {tecnica['nombre']}\n"
            f"Concepto cientifico: {tecnica['concepto_cientifico']}\n"
            f"Pasos:\n{pasos_texto}\n"
            f"Ejemplo aplicado al temario: {tecnica['ejemplo_aplicado']}"
        )
        documentos.append(
            Document(
                page_content=texto,
                metadata={"tipo": "tecnica_estudio", "clave": clave, "archivo": "TECNICAS_ESTUDIO_PRO"},
            )
        )

    return documentos


def _sanear_texto_unicode(texto: str) -> str:
    """Elimina surrogates UTF-16 sueltos que pypdf a veces deja al extraer
    texto de PDFs con fuentes/codificaciones no estandar (frecuente en PDFs
    escaneados/exportados desde Word). Sin esto, Chroma (chromadb, bindings
    en Rust) revienta con UnicodeEncodeError al indexar ese fragmento,
    tumbando el vectorstore entero (afecta al Tutor IA y a los Simulacros
    por igual, ya que comparten el mismo vectorstore)."""
    return texto.encode("utf-8", errors="ignore").decode("utf-8")


def _quitar_cabeceras_repetidas(documentos: list[Document]) -> None:
    """Quita, en el sitio, las lineas que se repiten identicas en la mayoria
    de "paginas" de un mismo archivo (cabeceras/pies de pagina tipo
    "PROCEDIMIENTO OPERATIVO ... Pagina X de 21 ..." que un PDF paginado
    repite en cada pagina).

    Sin esto, esas lineas cortas y repetidas comparten mucho vocabulario con
    el titulo/tema del propio documento, asi que la busqueda semantica las
    recupera antes que el contenido real (tablas, listados especificos) y
    los primeros resultados quedan copados de cabeceras vacias de contenido.
    Bug real encontrado con TEMA-36-PROCEDIMIENTOS-DE-TRENES-DE-SALIDA: una
    pregunta legitima sobre el Parque BC-2 (que si esta en el documento)
    recuperaba solo cabeceras de pagina repetidas, nunca la tabla de
    vehiculos de ese parque.

    Los TXT (un unico Document por archivo, sin paginacion) nunca disparan
    esto: el umbral exige al menos 3 "paginas" del mismo archivo.

    Las cabeceras/pies con numero de pagina ("Pagina 5 de 21...") no son
    identicas byte a byte entre paginas -- se comparan normalizando primero
    los digitos (\\d+ -> #), pero se eliminan las lineas originales, no la
    version normalizada."""
    por_archivo: dict[str, list[Document]] = {}
    for documento in documentos:
        por_archivo.setdefault(documento.metadata["archivo"], []).append(documento)

    for paginas in por_archivo.values():
        if len(paginas) < 3:
            continue

        conteo_normalizadas: dict[str, int] = {}
        for pagina in paginas:
            normalizadas_de_esta_pagina = {
                re.sub(r"\d+", "#", linea.strip())
                for linea in pagina.page_content.split("\n")
                if linea.strip()
            }
            for normalizada in normalizadas_de_esta_pagina:
                conteo_normalizadas[normalizada] = conteo_normalizadas.get(normalizada, 0) + 1

        umbral = max(3, len(paginas) * 0.5)
        normalizadas_repetidas = {
            normalizada for normalizada, veces in conteo_normalizadas.items()
            if veces >= umbral and len(normalizada) <= 200
        }
        if not normalizadas_repetidas:
            continue

        for pagina in paginas:
            pagina.page_content = "\n".join(
                linea for linea in pagina.page_content.split("\n")
                if re.sub(r"\d+", "#", linea.strip()) not in normalizadas_repetidas
            )


_PATRON_MARCADOR_PARQUE = re.compile(r"^PARQUE\b.*\bBC-\d+\b")


def _marcar_paginas_con_parque_actual(documentos: list[Document]) -> None:
    """Anota en la metadata de cada "pagina" (clave 'parque_actual') la
    ultima etiqueta de parque ("PARQUE MOVIL DE BC-2", "PARQUE DE LA SIERRA
    ORIENTAL, BC-2"...) vista en paginas anteriores del mismo archivo,
    mientras no aparezca una etiqueta nueva. No toca el contenido todavia
    -- eso lo hace _propagar_marcador_a_fragmentos(), DESPUES de trocear
    (ver su docstring para el motivo).

    Solo dispara en lineas que EMPIEZAN por "PARQUE" en mayusculas exactas
    (no cualquier mencion de "BC-N" en mitad de una frase, como las
    referencias cruzadas de TEMA-35 tipo "...a la que acude BC-9", ni un
    salto de linea de pypdf a mitad de frase en minusculas -- caso real
    encontrado en TEMA-34: "parque BC-2, se denomina FP-2." es la
    continuacion de una oracion, no una cabecera de seccion) y en archivos
    de al menos 2 "paginas" -- los TXT (una sola pagina) nunca lo disparan."""
    por_archivo: dict[str, list[Document]] = {}
    for documento in documentos:
        por_archivo.setdefault(documento.metadata["archivo"], []).append(documento)

    for paginas in por_archivo.values():
        if len(paginas) < 2:
            continue
        paginas_en_orden = sorted(paginas, key=lambda d: d.metadata.get("page", 0))

        marcador_actual = None
        for pagina in paginas_en_orden:
            for linea in pagina.page_content.split("\n"):
                linea_limpia = linea.strip()
                if len(linea_limpia) <= 80 and _PATRON_MARCADOR_PARQUE.match(linea_limpia):
                    marcador_actual = linea_limpia
            if marcador_actual:
                pagina.metadata["parque_actual"] = marcador_actual


def _propagar_marcador_a_fragmentos(fragmentos: list[Document]) -> None:
    """Antepone, en el sitio, la etiqueta de parque guardada en
    metadata['parque_actual'] (ver _marcar_paginas_con_parque_actual) al
    contenido de cada FRAGMENTO ya troceado que no la mencione ya.

    Hacerlo tras trocear -- no antes, sobre la pagina completa -- importa:
    una pagina de ~2000 caracteres se puede partir en 2-3 fragmentos de
    1000, y si la etiqueta solo se antepusiera una vez a la pagina entera,
    solo el primer fragmento resultante la conservaria, los demas la
    perderian en el troceo igual que sin el fix. Bug real (TEMA-36): con la
    etiqueta puesta solo a nivel de pagina, una pregunta sobre el Parque
    BC-2 seguia recuperando el fragmento de OTRO parque (BC-1) en vez del de
    BC-2, porque la etiqueta quedaba diluida en un fragmento de pagina
    completa en vez de presente en cada fragmento pequeño."""
    for fragmento in fragmentos:
        marcador = fragmento.metadata.get("parque_actual")
        if marcador and marcador not in fragmento.page_content:
            fragmento.page_content = f"[{marcador}]\n" + fragmento.page_content


def _iterar_bloques_docx(documento_docx):
    """Recorre el cuerpo de un .docx en orden real de aparicion, devolviendo
    parrafos y tablas intercalados tal como los veria un lector humano.
    python-docx no ofrece esto de serie: `documento.paragraphs` y
    `documento.tables` son dos listas separadas que pierden el orden y no
    dicen que tabla iba entre que parrafos -- hay que iterar
    `documento.element.body` directamente y distinguir cada hijo por su tag
    XML (`CT_P` = parrafo, `CT_Tbl` = tabla)."""
    for hijo in documento_docx.element.body.iterchildren():
        if isinstance(hijo, CT_P):
            yield DocxParagraph(hijo, documento_docx)
        elif isinstance(hijo, CT_Tbl):
            yield DocxTable(hijo, documento_docx)


def _texto_de_tabla_docx(tabla: DocxTable) -> str:
    """Aplana una tabla de Word a texto: una linea por fila, celdas unidas
    con ' | '. Sin esto, el contenido de las tablas (frecuente en los
    documentos de referencia de parques/riesgos que deposita el usuario) se
    perderia por completo -- `documento.paragraphs` no incluye texto de
    dentro de tablas."""
    lineas = []
    for fila in tabla.rows:
        celdas = [celda.text.strip() for celda in fila.cells]
        if any(celdas):
            lineas.append(" | ".join(celdas))
    return "\n".join(lineas)


def _cargar_documentos_docx() -> list[Document]:
    """Carga cada .docx de conocimiento/ como un unico Document (sin
    paginacion, igual que los TXT), preservando parrafos y tablas en orden.
    Sin este loader, cualquier .docx depositado en conocimiento/ se ignoraba
    en silencio -- ya paso de verdad con varios documentos reales del
    usuario (preguntas de repaso, contenido de la provincia de Huelva, base
    de conocimiento de parques de bomberos con 20 tablas)."""
    documentos = []
    for ruta in sorted(CARPETA_CONOCIMIENTO.glob("*.docx")):
        doc = DocxDocument(str(ruta))
        partes = []
        for bloque in _iterar_bloques_docx(doc):
            if isinstance(bloque, DocxParagraph):
                if bloque.text.strip():
                    partes.append(bloque.text)
            else:
                texto_tabla = _texto_de_tabla_docx(bloque)
                if texto_tabla:
                    partes.append(texto_tabla)
        texto = "\n".join(partes)
        if texto.strip():
            documentos.append(Document(page_content=texto, metadata={"source": str(ruta)}))
    return documentos


def _cargar_documentos_xlsx() -> list[Document]:
    """Carga cada .xlsx de conocimiento/ como un Document por hoja, uniendo
    las celdas no vacias de cada fila con ' | '. No asume una estructura de
    columnas fija: las hojas reales de este proyecto no siguen todas el
    mismo layout (ej. GENTILICIOS.xlsx repite pares municipio/gentilicio en
    columnas A-B, C-D y E-F de la misma fila, y su Hoja2 tiene datos de
    superficie en vez de gentilicio en la columna B -- error del propio
    archivo fuente, no de este loader). `data_only=True` para leer valores
    calculados de formulas en vez de la formula en si; `read_only=True`
    porque solo se necesita lectura, evita cargar estilos innecesarios."""
    documentos = []
    for ruta in sorted(CARPETA_CONOCIMIENTO.glob("*.xlsx")):
        libro = openpyxl.load_workbook(str(ruta), data_only=True, read_only=True)
        for hoja in libro.worksheets:
            lineas = []
            for fila in hoja.iter_rows(values_only=True):
                celdas = [str(valor).strip() for valor in fila if valor not in (None, "")]
                if celdas:
                    lineas.append(" | ".join(celdas))
            texto = "\n".join(lineas)
            if texto.strip():
                documentos.append(
                    Document(
                        page_content=f"{ruta.stem} - {hoja.title}\n{texto}",
                        metadata={"source": str(ruta)},
                    )
                )
    return documentos


def _tabla_a_markdown(tabla: list[list]) -> str:
    """Convierte una tabla extraida por pdfplumber (lista de filas, cada fila
    una lista de celdas) a una tabla Markdown (formato pipe). Necesario
    porque PyPDFDirectoryLoader (pypdf) no entiende la estructura de una
    tabla: solo extrae el texto en el orden en que aparece en el PDF, asi
    que una tabla de varias columnas se aplana a una lista lineal de
    numeros sin ninguna separacion visual entre filas/columnas. Bug real
    verificado (2026-08-16): con ese volcado plano, el tutor confundio la
    columna de 200 l/min con la de 1000 l/min en una tabla de perdida de
    carga de mangueras y dio una cifra incorrecta con total seguridad. Una
    tabla Markdown, en cambio, preserva la posicion de cada valor de forma
    explicita fila a fila."""
    filas = [
        [str(celda).strip() if celda is not None else "" for celda in fila]
        for fila in tabla
        if any(celda not in (None, "") for celda in fila)
    ]
    if len(filas) < 2 or len(filas[0]) < 2:
        return ""

    cabecera = filas[0]
    ancho = len(cabecera)
    lineas = [
        "| " + " | ".join(cabecera) + " |",
        "| " + " | ".join(["---"] * ancho) + " |",
    ]
    for fila in filas[1:]:
        fila = (fila + [""] * ancho)[:ancho]
        lineas.append("| " + " | ".join(fila) + " |")
    return "\n".join(lineas)


def _extraer_tablas_pdf(ruta: Path) -> list[Document]:
    """Extrae las tablas reales (con estructura de filas/columnas, via
    pdfplumber) de un PDF y las devuelve como Document adicionales en
    formato Markdown. Complementa, no sustituye, el texto plano que ya
    carga PyPDFDirectoryLoader para el mismo archivo -- asi no se toca la
    extraccion de texto ya verificada (cabeceras repetidas, marcador de
    parque en TEMA-36, etc.), solo se añade una version sin ambiguedad de
    cada tabla junto a ella."""
    documentos = []
    with pdfplumber.open(str(ruta)) as pdf:
        for numero_pagina, pagina in enumerate(pdf.pages):
            for tabla in pagina.extract_tables():
                markdown = _tabla_a_markdown(tabla)
                if markdown:
                    documentos.append(
                        Document(
                            page_content=f"Tabla (pagina {numero_pagina + 1}):\n{markdown}",
                            metadata={"source": str(ruta), "page": numero_pagina},
                        )
                    )
    return documentos


def _cargar_documentos_conocimiento() -> list[Document]:
    """Escanea backend/app/conocimiento/ y carga todos los PDF, TXT, DOCX y
    XLSX que encuentre."""
    documentos: list[Document] = []

    documentos.extend(PyPDFDirectoryLoader(str(CARPETA_CONOCIMIENTO)).load())

    for ruta_pdf in sorted(CARPETA_CONOCIMIENTO.glob("*.pdf")):
        documentos.extend(_extraer_tablas_pdf(ruta_pdf))

    cargador_txt = DirectoryLoader(
        str(CARPETA_CONOCIMIENTO),
        glob="**/*.txt",
        loader_cls=TextLoader,
        loader_kwargs={"encoding": "utf-8"},
        silent_errors=True,
    )
    documentos.extend(cargador_txt.load())

    documentos.extend(_cargar_documentos_docx())
    documentos.extend(_cargar_documentos_xlsx())

    for documento in documentos:
        documento.page_content = _sanear_texto_unicode(documento.page_content)
        # Nombre de archivo (no ruta completa: difiere entre Windows y el
        # servidor Linux) para poder filtrar el RAG de los Simulacros por
        # documento concreto (ver generar_banco.py, TEMA_A_ARCHIVOS) sin que
        # cada tema recupere por pura similitud semantica fragmentos de
        # documentos que no le corresponden.
        documento.metadata["archivo"] = Path(documento.metadata["source"]).name

    _quitar_cabeceras_repetidas(documentos)
    _marcar_paginas_con_parque_actual(documentos)

    return documentos


def _indice_persistido_existe() -> bool:
    """True si CARPETA_PERSISTENCIA_VECTORSTORE ya tiene un indice de Chroma
    creado en una ejecucion anterior (carpeta presente y con contenido)."""
    return CARPETA_PERSISTENCIA_VECTORSTORE.exists() and any(
        CARPETA_PERSISTENCIA_VECTORSTORE.iterdir()
    )


def _obtener_vectorstore() -> Chroma:
    """Devuelve el vectorstore compartido por el Tutor IA y los Simulacros,
    cargandolo una sola vez por proceso (_vectorstore como singleton lazy).

    CASO A (ya indexado): si CARPETA_PERSISTENCIA_VECTORSTORE tiene datos de
    una ejecucion anterior, se abre directamente desde disco (rapido, sin
    leer PDFs ni llamar a Gemini para generar embeddings).
    CASO B (primera vez / carpeta vacia o borrada): se reconstruye el indice
    completo desde los PDFs/TXT de conocimiento/ y se persiste en disco para
    que la proxima carga sea instantanea.
    """
    global _vectorstore
    if _vectorstore is not None:
        return _vectorstore

    embeddings = GoogleGenerativeAIEmbeddings(model=MODELO_EMBEDDINGS)

    if _indice_persistido_existe():
        _vectorstore = Chroma(
            persist_directory=str(CARPETA_PERSISTENCIA_VECTORSTORE),
            embedding_function=embeddings,
        )
        return _vectorstore

    documentos = _construir_documentos_diccionarios() + _cargar_documentos_conocimiento()
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP
    )
    fragmentos = splitter.split_documents(documentos)
    _propagar_marcador_a_fragmentos(fragmentos)
    _vectorstore = Chroma.from_documents(
        fragmentos,
        embeddings,
        persist_directory=str(CARPETA_PERSISTENCIA_VECTORSTORE),
    )
    return _vectorstore


def preguntar_al_tutor(query: str) -> str:
    """Busca los fragmentos mas relevantes del Plan Pro y genera la respuesta
    del tutor IA restringida a ese contexto."""
    vectorstore = _obtener_vectorstore()
    fragmentos = vectorstore.similarity_search(query, k=FRAGMENTOS_A_RECUPERAR)
    contexto = "\n\n---\n\n".join(
        f"[Fuente: {doc.metadata.get('archivo', 'desconocida')}]\n{doc.page_content}"
        for doc in fragmentos
    )

    llm = ChatGoogleGenerativeAI(model=MODELO_CHAT, temperature=0.3)
    mensajes = [
        SystemMessage(content=SYSTEM_PROMPT),
        HumanMessage(
            content=(
                f"Contexto disponible de tu Plan Pro:\n{contexto}\n\n"
                f"Pregunta del usuario: {query}"
            )
        ),
    ]
    respuesta = llm.invoke(mensajes)
    return respuesta.content


SYSTEM_PROMPT_PLAN_ESTUDIO = (
    "Eres un analista experto en oposiciones públicas. Tu objetivo es leer "
    "convocatorias oficiales y extraer un plan de estudio milimétrico."
)


def generar_plan_estudio_convocatoria(titulo_plaza: str, requisitos_minimos: str) -> str:
    """Genera un plan de ataque en Markdown para una convocatoria concreta del
    Tablon de Plazas Premium, adaptado exclusivamente al cuerpo real de la
    plaza (el prompt ya no fuerza temario de bomberos/hidraulica de forma
    fija: el propio texto le exige identificar el cuerpo antes de redactar,
    para que una plaza administrativa no reciba temario de bomberos)."""
    user_prompt = (
        "Analiza la siguiente convocatoria oficial. "
        f"Plaza: '{titulo_plaza}'. Extracto: '{requisitos_minimos}'. "
        "INSTRUCCIONES ESTRICTAS: "
        "1. Identifica el cuerpo exacto basándote en el título. "
        "2. Resume los requisitos reales o plazos. "
        "3. Construye un plan de ataque estructurado (títulos, subtítulos y "
        "listas con viñetas) adaptado EXCLUSIVAMENTE a la naturaleza de esta "
        "plaza (ej. si es puramente administrativa, enfócate en legislación; "
        "si es de bomberos, en fuego e hidráulica). NO mezcles temarios. "
        "4. Si el extracto está muy vacío, avisa al usuario de consultar las "
        "bases, pero ofrécele las materias troncales habituales del cuerpo.\n\n"
        "FORMATO DE LA RESPUESTA (muy importante): toda la respuesta, "
        "incluido el plan de ataque, debe escribirse directamente en "
        "Markdown normal (# para títulos, ** para negrita, - para listas). "
        "NUNCA envuelvas el plan de ataque ni ninguna otra parte de la "
        "respuesta dentro de un bloque de código de tres comillas invertidas "
        "(```), ni la trates como si fuera un fragmento de código o una "
        "plantilla aparte: es texto normal con formato Markdown, se renderiza "
        "directamente en la pantalla del usuario."
    )
    llm = ChatGoogleGenerativeAI(model=MODELO_CHAT, temperature=0.3)
    mensajes = [
        SystemMessage(content=SYSTEM_PROMPT_PLAN_ESTUDIO),
        HumanMessage(content=user_prompt),
    ]
    respuesta = llm.invoke(mensajes)
    return respuesta.content


